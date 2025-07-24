#!/usr/bin/env python3
"""
Document Processing Test Script

This script tests the document processing capabilities for the RAG system,
including support for different file formats and text extraction.
"""
import os
import sys
import time
from pathlib import Path
from typing import Dict, List, Any, Optional
import traceback

# Add the project root to the path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

def test_markdown_processing() -> Dict[str, Any]:
    """Test markdown document processing."""
    print("📝 Testing markdown processing...")
    
    try:
        import markdown
        
        # Test with sample markdown content
        sample_md = """
# Test Document

## Section 1
This is a test section with some content.

### Subsection 1.1
- Point 1
- Point 2
- Point 3

## Section 2
More content here with **bold** and *italic* text.

### Code Example
```python
def test_function():
    return "Hello World"
```
"""
        
        # Convert to HTML
        html = markdown.markdown(sample_md)
        
        # Extract text content (simple approach)
        import re
        text_content = re.sub(r'<[^>]+>', '', html)
        text_content = re.sub(r'\s+', ' ', text_content).strip()
        
        return {
            'status': 'success',
            'original_length': len(sample_md),
            'processed_length': len(text_content),
            'html_generated': len(html) > 0,
            'text_extracted': len(text_content) > 0
        }
        
    except Exception as e:
        return {
            'status': 'error',
            'error': str(e)
        }

def test_pdf_processing() -> Dict[str, Any]:
    """Test PDF document processing."""
    print("📄 Testing PDF processing...")
    
    try:
        import pypdf
        
        # Create a simple test PDF
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        test_pdf_path = "test_document.pdf"
        
        # Create test PDF
        c = canvas.Canvas(test_pdf_path, pagesize=letter)
        c.drawString(100, 750, "Test PDF Document")
        c.drawString(100, 700, "This is a test document for RAG system processing.")
        c.drawString(100, 650, "It contains multiple lines of text.")
        c.drawString(100, 600, "This will be used to test PDF extraction capabilities.")
        c.save()
        
        # Extract text from PDF
        with open(test_pdf_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            text_content = ""
            
            for page in pdf_reader.pages:
                text_content += page.extract_text()
        
        # Clean up test file
        os.remove(test_pdf_path)
        
        return {
            'status': 'success',
            'pages_processed': len(pdf_reader.pages),
            'text_extracted': len(text_content) > 0,
            'extracted_length': len(text_content)
        }
        
    except Exception as e:
        return {
            'status': 'error',
            'error': str(e)
        }

def test_docx_processing() -> Dict[str, Any]:
    """Test DOCX document processing."""
    print("📄 Testing DOCX processing...")
    
    try:
        from docx import Document
        
        # Create a simple test DOCX
        doc = Document()
        doc.add_heading('Test DOCX Document', 0)
        doc.add_paragraph('This is a test document for RAG system processing.')
        doc.add_heading('Section 1', level=1)
        doc.add_paragraph('This section contains important information.')
        doc.add_paragraph('It has multiple paragraphs with different content.')
        
        test_docx_path = "test_document.docx"
        doc.save(test_docx_path)
        
        # Extract text from DOCX
        doc = Document(test_docx_path)
        text_content = ""
        
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Clean up test file
        os.remove(test_docx_path)
        
        return {
            'status': 'success',
            'paragraphs_processed': len(doc.paragraphs),
            'text_extracted': len(text_content) > 0,
            'extracted_length': len(text_content)
        }
        
    except Exception as e:
        return {
            'status': 'error',
            'error': str(e)
        }

def test_text_chunking() -> Dict[str, Any]:
    """Test text chunking capabilities."""
    print("✂️  Testing text chunking...")
    
    try:
        # Sample text for chunking
        sample_text = """
# Technical Documentation

## Introduction
This is a comprehensive technical document that contains multiple sections with detailed information about various topics.

## Section 1: Overview
This section provides an overview of the system architecture and design principles. It includes information about the main components and their interactions.

### Subsection 1.1: Components
The system consists of several key components:
- Database layer for data persistence
- API layer for external communication
- Business logic layer for processing
- User interface layer for interaction

## Section 2: Implementation
This section details the implementation approach and coding standards.

### Subsection 2.1: Coding Standards
All code must follow established standards:
- Use meaningful variable names
- Include proper documentation
- Follow the DRY principle
- Implement proper error handling

## Section 3: Testing
This section covers testing strategies and methodologies.

### Subsection 3.1: Unit Testing
Unit tests are essential for code quality:
- Test individual functions
- Achieve high code coverage
- Use meaningful test names
- Mock external dependencies
"""
        
        # Simple chunking by sections
        chunks = []
        current_chunk = ""
        lines = sample_text.split('\n')
        
        for line in lines:
            if line.startswith('##') and current_chunk:
                # Start new chunk
                chunks.append(current_chunk.strip())
                current_chunk = line + '\n'
            else:
                current_chunk += line + '\n'
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return {
            'status': 'success',
            'original_length': len(sample_text),
            'chunks_created': len(chunks),
            'average_chunk_length': sum(len(chunk) for chunk in chunks) / len(chunks) if chunks else 0,
            'chunk_sizes': [len(chunk) for chunk in chunks]
        }
        
    except Exception as e:
        return {
            'status': 'error',
            'error': str(e)
        }

def test_embedding_generation() -> Dict[str, Any]:
    """Test embedding generation for document chunks."""
    print("🧠 Testing embedding generation...")
    
    try:
        from src.core.chroma_embedding_manager_simple import SimpleChromaEmbeddingManager
        
        chroma_manager = SimpleChromaEmbeddingManager()
        
        # Test chunks
        test_chunks = [
            "This is a test chunk about technical documentation.",
            "Another chunk about project management processes.",
            "A third chunk about coding standards and best practices."
        ]
        
        start_time = time.time()
        embeddings = chroma_manager.get_batch_embeddings(test_chunks)
        generation_time = time.time() - start_time
        
        return {
            'status': 'success',
            'chunks_processed': len(test_chunks),
            'embeddings_generated': len(embeddings),
            'generation_time': generation_time,
            'average_time_per_chunk': generation_time / len(test_chunks) if test_chunks else 0,
            'embedding_dimension': len(list(embeddings.values())[0]) if embeddings else 0
        }
        
    except Exception as e:
        return {
            'status': 'error',
            'error': str(e)
        }

def run_all_tests() -> Dict[str, Any]:
    """Run all document processing tests."""
    print("🚀 Running document processing tests...\n")
    
    tests = {
        'markdown': test_markdown_processing(),
        'pdf': test_pdf_processing(),
        'docx': test_docx_processing(),
        'chunking': test_text_chunking(),
        'embeddings': test_embedding_generation()
    }
    
    # Calculate overall success rate
    successful_tests = sum(1 for result in tests.values() if result['status'] == 'success')
    total_tests = len(tests)
    success_rate = (successful_tests / total_tests) * 100 if total_tests > 0 else 0
    
    return {
        'tests': tests,
        'successful_tests': successful_tests,
        'total_tests': total_tests,
        'success_rate': success_rate,
        'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
    }

def print_test_results(results: Dict[str, Any]):
    """Print formatted test results."""
    print("\n" + "="*80)
    print("📋 DOCUMENT PROCESSING TEST RESULTS")
    print("="*80)
    print(f"Timestamp: {results['timestamp']}")
    print(f"Success Rate: {results['success_rate']:.1f}% ({results['successful_tests']}/{results['total_tests']})")
    print("="*80)
    
    for test_name, result in results['tests'].items():
        print(f"\n🔍 {test_name.upper()} TEST:")
        if result['status'] == 'success':
            print("✅ PASSED")
            for key, value in result.items():
                if key != 'status':
                    print(f"   {key}: {value}")
        else:
            print("❌ FAILED")
            print(f"   Error: {result.get('error', 'Unknown error')}")
    
    print("\n" + "="*80)
    
    if results['success_rate'] >= 80:
        print("✅ Document processing is ready for RAG implementation")
    elif results['success_rate'] >= 60:
        print("⚠️  Some issues detected, but processing is mostly functional")
    else:
        print("❌ Significant issues detected, review required")

def main():
    """Main test function."""
    try:
        results = run_all_tests()
        print_test_results(results)
        
        # Return exit code based on success rate
        if results['success_rate'] >= 80:
            return 0
        else:
            return 1
            
    except Exception as e:
        print(f"\n❌ Test execution failed: {e}")
        print(traceback.format_exc())
        return 1

if __name__ == "__main__":
    exit(main()) 